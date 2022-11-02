# -*- coding: utf-8 -*-
"""
Created on Sun Nov 21 11:14:14 2021

@author: shenhong
"""

from docx import Document
#import docx
import win32com.client as wc
import os


#获取文件夹下的所有doc文档
'''
path="C:\\Users\\Wyc\\Desktop\\WQP\\000"

files= os.listdir(path)
filenames=[]

for file in files:
    if file.endswith(".doc"):
        filenames.append(file)
        
        
        
        




#doc文件另存为docx


word = wc.Dispatch("Word.Application")
for file in filenames:
    filepath=path+"\\"+file
    namenew=path+"\\"+file+"x"

    doc = word.Documents.Open(r"C:\\Users\\Wyc\\Desktop\\WQP\\000\\"+file)
    #上面的地方只能使用完整绝对地址，相对地址找不到文件，且，只能用“\\”，不能用“/”，哪怕加了 r 也不行，涉及到将反斜杠看成转义字符。
    doc.SaveAs(r"C:\\Users\\Wyc\\Desktop\\WQP\\000\\"+file+"x", 12, False, "", True, "", False, False, False, False)#转换后的文件,12代表转换后为docx文件
    #doc.SaveAs(r"F:\\***\\***\\appendDoc\\***.docx", 12)#或直接简写
    #注意SaveAs会打开保存后的文件，有时可能看不到，但后台一定是打开的
doc.Close
word.Quit
'''

#获取文件夹下的所有docx文档

path="C:\\Users\\Wyc\\Desktop\\WQP\\PDF"

files= os.listdir(path)
filenamedocx=[]

for file in files:
    if file[0] != "~":
        
        if file.endswith(".docx"):
            filenamedocx.append(file)





timeall=[]

pointnameall=[]
PHall=[]
CODall=[]
DOall=[]
NH3_Nall=[]
WaterQuailityall=[]

for filename in filenamedocx:
    

    doc = Document(filename)
'''
    tb1=doc.tables[0]
    row=tb1.rows
    
    
    timed=doc.paragraphs[4].text
    timed=timed.split("，")
    timed=timed[0]

    timeall.append(timed)
    pointnameall.append(timed)
    PHall.append(timed)
    CODall.append(timed)
    DOall.append(timed)
    NH3_Nall.append(timed)
    WaterQuailityall.append(timed)
    
  
    
    print(timed)
    for r in tb1.rows:
        row_cells=r.cells
        bianhao=row_cells[0].text
        pname=row_cells[3].text
        #print(pname)
        PH=row_cells[5].text
        DO=row_cells[6].text
        COD=row_cells[7].text
        NH3_N=row_cells[8].text
        WaterQuaility=row_cells[9].text
        #print(WaterQuaility)
        pointnameall.append(pname)
        PHall.append(PH)
        CODall.append(COD)
        DOall.append(DO)
        NH3_Nall.append(NH3_N)
        WaterQuailityall.append(WaterQuaility)
        
        
#列表输出为excel

result = open('resultdata.xls', 'w', encoding='gbk')
#result.write('X\tY\n')
for m in range(len(pointnameall)):
    result.write(str(pointnameall[m]))
    result.write('\t')
    
    result.write(str(PHall[m]))
    result.write('\t')
    
    result.write(str(DOall[m]))
    result.write('\t')
    
    result.write(str(CODall[m]))
    result.write('\t')
    
    result.write(str(NH3_Nall[m]))
    result.write('\t')
    
    result.write(str(WaterQuailityall[m]))
    result.write('\t')
    
    result.write('\n')
    
    
    
result.close()
'''
    


        
        
        