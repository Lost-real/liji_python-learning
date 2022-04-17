from docx import Document

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn

import time
price =input('请输入工资调整金额：')
name01=open('name.txt',"r",encoding="utf-8")
# for i in name01:
#     i=i.split()
#     print(i[0])
# name02=list(name01)
# # print(name02)
# dict={}
# dict[name02[0]]=name02[1]
# # namelist02=list(name)
# print(dict)
# company_list=['小明','小李','小红','小兰','小王','小丽','小美']
#今天的时间
today=time.strftime("%Y{y}%m{m}%d{d}",time.localtime()).format(y="年",m="月",d="日")
for i in name01:
    i=i.split()
# for i in company_list:
    document=Document()
    #设置文档的基础字体
    document.styles['Normal'].font.name=u'宋体'
    #识别中文字体
    document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'),u'宋体')
    #建立一个自然段
    p1=document.add_paragraph()
    #对齐方式剧中，没有这句话的默认左对齐
    p1.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run1=p1.add_run("关于%s工资待遇调整的通知"%(today))
    run1.font.name="微软雅黑"
    run1.element.rPr.rFonts.set(qn('w:eastAsia'),u'微软雅黑')
    run1.font.size=Pt(21)
    run1.font.bold=True
    p1.space_after=Pt(5)
    p1.space_before = Pt(5)

    p2=document.add_paragraph()
    run2=p2.add_run(i[0]+':')
    run2.font.name = "宋体"
    run2.element.rPr.rFonts.set(qn('w:eastAsia'),u'宋体')
    run2.font.size=Pt(16)
    run2.font.bold=True

    p3 = document.add_paragraph()
    run3 = p3.add_run("ID"+i[1]+"因为受疫情的影响，我们很抱歉通知您，您的工资调整为每月%s元，特此通知" % price)
    run3.font.name = "宋体"
    run3.element.rPr.rFonts.set(qn('w:eastAsia'),u'宋体')
    run3.font.size=Pt(14)

    p4 = document.add_paragraph()
    p4.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    run4 = p4.add_run("人事部：李经理 电话 88888888")
    run4.font.name = "宋体"
    run4.element.rPr.rFonts.set(qn('w:eastAsia'),u'宋体')
    run4.font.size=Pt(14)
    run4.font.bold=True
#注意，如果是i[0]，则保存的文件名是"小红-工资待遇调整通知.docx"；如果是i，则保存的文件名是"['小红', '123']-工资待遇调整通知.docx"
    document.save('%s-工资待遇调整通知.docx'%i[0])