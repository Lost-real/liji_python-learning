#coding:utf-8

import os
#注意win32com的包的名字是pywin32
import win32com
from win32com.client import Dispatch, constants
from docx import Document
b=open("E:\\pythonworking\\file\\aa1\\a.txt","w")

def parse_docx(f):
  """读取docx，返回姓名和行业
  """
  d = Document(f)#读入docx文件
  t = d.tables[0]#读入第一个表格
  par1=d.paragraphs[4].text#读入第五段
  # par2=par1.runs[1]
  par2=par1.split(":")[1]#以第五段中的冒号":"为分隔符
  # par3=par2[1]

  # for i in d.paragraphs:
  #   print(i.text)
  # t_next=d.tables[1]
  #title= t.cell(0,1).text//比如说标题在表格中位于（0，1）
  在编教职工数 = t.cell(1,5).text#表示在编教职工数位于表格中第一行第五列，注意这个需要试，还有，行列的起始位置都是从0开始计算的
  实际教职工数 = t.cell(1,10).text #1
  一五年在校生数 = t.cell(6,10).text   #1
  班级数=t.cell(8,10).text   #1

#   print(par2+":"+在编教职工数+":"+实际教职工数+":"+一五年在校生数+":"+班级数,end='\n ')
#b.write()的意思是保存输出的结果，其中os.path.splitext(doc)[0]是docx的名字，但是不包含.docx
#   b.write(os.path.splitext(doc)[0]+":"+par2+":"+在编教职工数+":"+实际教职工数+":"+一五年在校生数+":"+班级数+"\n")


''' 上述函数主要实现文件的读取 '''
if __name__ == "__main__":
  w = win32com.client.Dispatch('Word.Application')

  # 遍历文件
  PATH = "E:\\pythonworking\\file\\aa1" # windows文件路径
  doc_files = os.listdir(PATH)
  for doc in doc_files:
    if os.path.splitext(doc)[1] == '.docx':
      # print(os.path.splitext(doc)[0])
      try:
        parse_docx(PATH+'\\'+doc)
      except Exception as e:
        print(e)


  # def write_excel(workbook, i_in, par2_in, 在编教职工数_in, 实际教职工数_in, 一五年在校生数_in, 班级数_in):
  #
  #   first_sheet = workbook.Worksheets(1) # 需要获取excel
  #   first_sheet.Cells(i_in, 3).value = par2_in
  #   first_sheet.Cells(i_in, 4).value = 在编教职工数_in
  #   first_sheet.Cells(i_in, 5).value = 实际教职工数_in
  #   first_sheet.Cells(i_in, 6).value = 一五年在校生数_in
  #   first_sheet.Cells(i_in, 7).value = 班级数_in
  #   # first_sheet.Cells(i_in, 8).value = Birth_in
  #   # first_sheet.Cells(i_in, 9).value = Telephone_in
  #   # first_sheet.Cells(i_in, 10).value = QQ_in
  #   # first_sheet.Cells(i_in, 11).value = QQEmail_in
  #   # first_sheet.Cells(i_in, 12).value = Place_in
  #   # first_sheet.Cells(i_in, 13).value = Address_in
  #   #
  #   # first_sheet.Cells(i_in, 14).value = Weeks_in
  #   # first_sheet.Cells(i_in, 15).value = participate_in
  #   # first_sheet.Cells(i_in, 16).value = other_team_in
  #   # first_sheet.Cells(i_in, 2).value = Data_number
  #   # first_sheet.Cells(i_in, 1).value = Witch_team_in
  #
  #   print('成功写入:' + name_in + " 的信息", "这是第", i_in - rember_ever + 1, "个")
  #   workbook.Save()