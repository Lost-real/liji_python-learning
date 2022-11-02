#coding:utf-8
import win32com.client
import os
import win32com

from win32com.client import Dispatch, constants
from docx import Document
import time

def parse_docx(f,i,f_excel,Data_number):
  """读取docx，返回姓名和行业
  """
  d = Document(f)
  t = d.tables[0]
  t_next=d.tables[1]
  #title= t.cell(0,1).text
  name = t.cell(1,1).text     #1
  Gender = t.cell(1,3).text   #1
  Race = t.cell(1,5).text     #1
  School=t.cell(2,1).text     #1
  Edu_back=t.cell(2,4).text   #1
  Major=t.cell(3,1).text      #1
  Grade=t.cell(3,4).text      #1
  Poli_Status=t.cell(4,1).text#1
  Place=t.cell(4,3).text      #1
  ID_card=t.cell(5,1).text    #1
  Date_birth=t.cell(5,4).text #1
  Telephone=t.cell(6,1).text  #1
  E_mail=t.cell(6,4).text     #1
  QQ=t.cell(7,1).text         #1
  Address=t.cell(7,4).text    #1
  Which_team=t.cell(8,1).text #1
  Weeks=t.cell(9,4).text
  participate=t.cell(10,4).text
  other_team=t.cell(11,4).text

#  Other=t_next.cell(1,1).text
  print(name, Gender,Race,School,Edu_back,Major,Grade,Poli_Status,Place,ID_card,Date_birth,Telephone,E_mail,QQ,Place,Address,Which_team,
        Weeks,participate,other_team,end='\n ')
  write_excel(f_excel,i,name,Gender,School,Major,Grade,Date_birth,Telephone, E_mail,QQ,Place,Address,Data_number,
              Weeks,participate,other_team,Which_team)
  rename("E:\\pythonworking\\file\word001",f,name,Data_number)

def write_excel(workbook,i_in,name_in,Gender_in,School_in,Major_in,Grade_in,Birth_in,Telephone_in,QQ_in,
                QQEmail_in,Place_in,Address_in,Data_number,Weeks_in,participate_in,other_team_in,Witch_team_in):

    first_sheet=workbook.Worksheets(1)
    first_sheet.Cells(i_in,3).value=name_in
    first_sheet.Cells(i_in,4).value=Gender_in
    first_sheet.Cells(i_in,5).value=School_in
    first_sheet.Cells(i_in,6).value=Major_in
    first_sheet.Cells(i_in,7).value=Grade_in
    first_sheet.Cells(i_in,8).value=Birth_in
    first_sheet.Cells(i_in,9).value=Telephone_in
    first_sheet.Cells(i_in,10).value=QQ_in
    first_sheet.Cells(i_in, 11).value = QQEmail_in
    first_sheet.Cells(i_in,12).value =Place_in
    first_sheet.Cells(i_in,13).value=Address_in

    first_sheet.Cells(i_in, 14).value = Weeks_in
    first_sheet.Cells(i_in, 15).value = participate_in
    first_sheet.Cells(i_in, 16).value = other_team_in
    first_sheet.Cells(i_in, 2).value = Data_number
    first_sheet.Cells(i_in, 1).value = Witch_team_in

    print('成功写入:'+name_in+" 的信息","这是第",i_in-rember_ever+1,"个")
    workbook.Save()

def rename(path,old_file,newname,Data_number):
    '''path="这里替换为你的文件夹的路径";
    filelist=os.listdir(path)#该文件夹下所有的文件（包括文件夹）
    for files in filelist:#遍历所有文件
    Olddir=os.path.join(path,files);#原来的文件路径
    if os.path.isdir(Olddir):#如果是文件夹则跳过
    continue;'''
    Olddir=os.path.join(path,old_file);#原来的文件路径
    filename=os.path.splitext(old_file)[0];#文件名
    filetype=os.path.splitext(old_file)[1];#文件扩展名
 #   data_identi=
    Newdir=os.path.join(path,str(Data_number)+newname+filetype);#新的文件路径
    os.rename(Olddir,Newdir);#重命名

def introduction():
    print("美仔你好，感谢使用本软件,它将会快速的帮你完成工作")
    print("开始前，请将今天你收集的所有报名表放到C:\My_Work文件夹下,如果已经成功放置，请回复1,之后按回车")
    rece_ancer=input()
    while  (rece_ancer!='1'):
             rece_ancer=input("还没准备好嘛，确认好了以后输入1\n")
    print("好的，五秒钟后开始工作，在这期间，请不要打开word文档和excel文档哦\n")
    time.sleep(1)
    print("5")
    time.sleep(1)
    print("4")
    time.sleep(1)
    print("3")
    time.sleep(1)
    print("2")
    time.sleep(1)
    print("1")


if __name__ == "__main__":

  data_identifier = int(input("请输入今天是几月几号，格式如：409、413、421,按回车结束\n"))
  while True:
        if(data_identifier>=101 and data_identifier<=1230):
             break
        else:
            print("请重新输入")
            data_identifier=int(input())
  data_identifier = data_identifier * 100+1
  print("好的，接下来请按提示操作",data_identifier)
  introduction()

  w = win32com.client.Dispatch('Word.Application')
  excel = win32com.client.Dispatch('Excel.Application')
  workbook=excel.Workbooks.open('E:\\pythonworking\\file\\word001\\海口分队报名信息录入表(总表)+日期.xlsx')
  excel.Visible=False

  i=3
  # 遍历文件
  PATH = "E:\\pythonworking\\file\\word001" # windows文件路径
  doc_files = os.listdir(PATH)

  rember_sheet=workbook.Worksheets(1)
  for rember_ever in range(3,1000):
          Value=rember_sheet.Cells(rember_ever,2).value
          print(Value,rember_ever)
          if(Value==None):
             break
  i=rember_ever
  print("从上次的第",i,"行录入\n")

  for doc in doc_files:
    if os.path.splitext(doc)[1] == '.docx':
      try:
        parse_docx(PATH+'\\'+doc,i,workbook,data_identifier)
        i=i+1
        data_identifier=data_identifier+1
      except Exception as e:
        print(e)
    elif os.path.splitext(doc)[1] == '.doc':
      try:
        parse_doc(PATH+'\\'+doc)
      except Exception as e:
        print (e)
  print("全部操作完成,好了，现在你可以到C:\My_Work文件夹下查看数据,","今天一共写了",i-rember_ever,"个人的信息\n")
  print("辛苦啦,现在请把C:\My_Work下的报名表复制到对应日期的文件夹内，并从“海口分队报名信息录入表(总表)+日期”这个文件里的信息\n复制到总的信息表上，压缩传到群里\n")
  excel.Application.Quit()
  time.sleep(150)